import os
from flask import Blueprint, render_template, session, url_for, flash, redirect, request, abort
from flask_login import login_required, current_user
from app import db
from app.models.student import Student
from app.models.class_model import ClassModel as Class
from app.forms import StudentForm, SearchStudentForm, ScoreForm
from app.decorators import admin_required
from sqlalchemy import or_
from flask import make_response
import io
from datetime import datetime
import pandas as pd
from docx import Document
from docx.shared import Inches
from io import BytesIO
from werkzeug.utils import secure_filename

student_bp = Blueprint('student', __name__)

@student_bp.route("/students", methods=['GET', 'POST'])
@login_required
def list_students():
    """Hiển thị danh sách thiếu nhi, có tìm kiếm và lọc."""
    search_form = SearchStudentForm(request.form if request.method == 'POST' else request.args)

    # Lấy danh sách lớp để đưa vào dropdown lọc
    if current_user.is_admin():
        classes = Class.query.order_by(Class.name).all()
        search_form.class_filter.choices = [('', 'Tất cả các lớp')] + [(str(c.id), c.name) for c in classes]
    elif current_user.role == 'guest':
        # Cho guest: có thể lọc theo tất cả các lớp
        classes = Class.query.order_by(Class.name).all()
        search_form.class_filter.choices = [('', 'Tất cả các lớp')] + [(str(c.id), c.name) for c in classes]
    else:
        # Cho leader: chỉ các lớp được phân công
        assigned_class_ids = [c.id for c in current_user.assigned_classes.all()]
        classes = Class.query.filter(Class.id.in_(assigned_class_ids)).order_by(Class.name).all()
        search_form.class_filter.choices = [('', 'Tất cả các lớp')] + [(str(c.id), c.name) for c in classes]

    # Bắt đầu query, join với Class để có thể sắp xếp/lọc theo tên lớp
    query = Student.query.join(Class).order_by(Class.name, Student.full_name)

    # Nếu không phải admin và không phải guest, chỉ lấy học sinh trong các lớp được phân công
    if not current_user.is_admin() and current_user.role != 'guest':
        assigned_class_ids = [c.id for c in current_user.assigned_classes.all()]
        query = query.filter(Student.class_id.in_(assigned_class_ids))

    search_term = None
    class_id_filter = None

    # Xử lý tìm kiếm và lọc khi submit form hoặc từ URL
    if request.method == 'POST' and 'submit_search' in request.form and search_form.validate():
         search_term = search_form.search_term.data
         class_id_filter = search_form.class_filter.data
    elif request.method == 'GET':
        search_term = request.args.get('search_term', None)
        class_id_filter = request.args.get('class_filter', None)
        # Cập nhật giá trị mặc định cho form từ URL
        if search_term:
             search_form.search_term.data = search_term
        if class_id_filter:
             search_form.class_filter.data = class_id_filter

    # Áp dụng bộ lọc tìm kiếm (nếu có)
    if search_term:
        search_pattern = f"%{search_term}%"
        query = query.filter(or_(
            Student.full_name.ilike(search_pattern),
            Student.ten_thanh.ilike(search_pattern)
        ))

    # Áp dụng bộ lọc lớp (nếu có)
    if class_id_filter:
        try:
            query = query.filter(Student.class_id == int(class_id_filter))
        except ValueError:
            pass # Bỏ qua nếu class_id không phải số nguyên

    students = query.all()

    return render_template('students.html',
                           students=students,
                           title="Danh sách thiếu nhi",
                           search_form=search_form) # <-- Truyền form tìm kiếm ra template


@student_bp.route("/student/add", methods=['GET', 'POST'])
@login_required
def add_student(class_id=None):
    """Thêm một thiếu nhi mới."""
    # Guest không được thêm thiếu nhi
    if current_user.role == 'guest':
        flash('Bạn không có quyền thêm thiếu nhi.', 'danger')
        abort(403)
    # Nếu có class_id, kiểm tra quyền
    if class_id:
        # Áp dụng decorator logic manually since we can't decorate here
        if not current_user.is_authenticated:
            abort(401)
        if not current_user.is_admin():
            is_assigned = current_user.assigned_classes.filter_by(id=class_id).first()
            if not is_assigned:
                flash('Bạn không có quyền thêm thiếu nhi vào lớp này.', 'danger')
                abort(403)

    form = StudentForm()
    # Kiểm tra xem có lớp nào chưa, nếu chưa thì không cho thêm SV
    if not Class.query.first():
         flash('Chưa có lớp học nào được tạo. Vui lòng tạo lớp trước khi thêm thiếu nhi.', 'warning')
         return redirect(url_for('class_admin.list_classes')) # Chuyển đến trang quản lý lớp

    # Nếu class_id được cung cấp, pre-set lớp trong form
    if class_id:
        class_obj = Class.query.get_or_404(class_id)
        form.class_obj.data = class_obj
    elif not current_user.is_admin():
        # Cho leader: chỉ hiển thị các lớp được phân công
        assigned_classes = current_user.assigned_classes.all()
        if assigned_classes:
            # Giữ lại query_factory nhưng filter chỉ các lớp được phân công
            form.class_obj.query_factory = lambda: assigned_classes
            # Mặc định chọn lớp đầu tiên
            form.class_obj.data = assigned_classes[0]

    if form.validate_on_submit():
        student = Student(
            full_name=form.full_name.data,
            ten_thanh=form.ten_thanh.data,
            date_of_birth=form.date_of_birth.data,
            gender=form.gender.data,
            ho_ten_bo=form.ho_ten_bo.data,
            ho_ten_me=form.ho_ten_me.data,
            sdt_phu_huynh=form.sdt_phu_huynh.data,
            # Lấy class_id từ đối tượng Class được chọn trong dropdown
            class_id=form.class_obj.data.id if form.class_obj.data else None,
            # Điểm số
            diem_mieng=form.diem_mieng.data,
            diem_giua_ki_1=form.diem_giua_ki_1.data,
            diem_cuoi_ki_1=form.diem_cuoi_ki_1.data,
            diem_giua_ki_2=form.diem_giua_ki_2.data,
            diem_cuoi_ki_2=form.diem_cuoi_ki_2.data
        )
        db.session.add(student)
        db.session.commit()
        flash('Đã thêm thiếu nhi mới thành công!', 'success')
        if class_id:
            return redirect(url_for('main_routes.view_class', class_id=class_id))
        return redirect(url_for('student.list_students'))
    return render_template('student_form.html', title='Thêm thiếu nhi', form=form, legend='Thêm thiếu nhi')

@student_bp.route("/student/<int:student_id>/update", methods=['GET', 'POST'])
@login_required
def update_student(student_id):
    """Cập nhật thông tin thiếu nhi."""
    student = Student.query.get_or_404(student_id)

    # Kiểm tra quyền: Admin hoặc Huynh trưởng của lớp, Guest không được sửa
    if current_user.role == 'guest':
        flash('Bạn không có quyền sửa thông tin thiếu nhi.', 'danger')
        abort(403)
    if not current_user.is_admin():
        if not student.class_id:
            flash('Thiếu nhi chưa được xếp lớp.', 'danger')
            abort(403)
        is_assigned = current_user.assigned_classes.filter_by(id=student.class_id).first()
        if not is_assigned:
            flash('Bạn không có quyền sửa thiếu nhi này.', 'danger')
            abort(403)

    form = StudentForm(obj=student)

    # Gán sẵn lớp hiện tại cho dropdown khi GET request
    if request.method == 'GET':
        form.class_obj.data = student.class_assigned

    if form.validate_on_submit():
        student.full_name = form.full_name.data
        student.ten_thanh = form.ten_thanh.data
        student.date_of_birth = form.date_of_birth.data
        student.gender = form.gender.data
        student.ho_ten_bo = form.ho_ten_bo.data
        student.ho_ten_me = form.ho_ten_me.data
        student.sdt_phu_huynh = form.sdt_phu_huynh.data
        # Lấy class_id từ đối tượng Class được chọn trong dropdown
        student.class_id = form.class_obj.data.id if form.class_obj.data else None
        # Điểm số
        student.diem_mieng = form.diem_mieng.data
        student.diem_giua_ki_1 = form.diem_giua_ki_1.data
        student.diem_cuoi_ki_1 = form.diem_cuoi_ki_1.data
        student.diem_giua_ki_2 = form.diem_giua_ki_2.data
        student.diem_cuoi_ki_2 = form.diem_cuoi_ki_2.data
        db.session.commit()
        flash('Thông tin thiếu nhi đã được cập nhật!', 'success')
        if current_user.is_admin():
            return redirect(url_for('student.list_students'))
        else:
            return redirect(url_for('main_routes.view_class', class_id=student.class_id))

    return render_template('student_form.html', title='Cập nhật thiếu nhi', form=form, legend='Cập nhật thiếu nhi')

@student_bp.route("/student/<int:student_id>/delete", methods=['POST'])
@login_required
def delete_student(student_id):
    """Xóa thiếu nhi."""
    student = Student.query.get_or_404(student_id)

    # Kiểm tra quyền: Admin hoặc Huynh trưởng của lớp, Guest không được xóa
    if current_user.role == 'guest':
        flash('Bạn không có quyền xóa thiếu nhi.', 'danger')
        abort(403)
    if not current_user.is_admin():
        if not student.class_id:
            flash('Thiếu nhi chưa được xếp lớp.', 'danger')
            abort(403)
        is_assigned = current_user.assigned_classes.filter_by(id=student.class_id).first()
        if not is_assigned:
            flash('Bạn không có quyền xóa thiếu nhi này.', 'danger')
            abort(403)

    db.session.delete(student)
    db.session.commit()
    flash('Thiếu nhi đã được xóa!', 'success')
    if current_user.is_admin():
        return redirect(url_for('student.list_students'))
    else:
        return redirect(url_for('main_routes.view_class', class_id=student.class_id))

@student_bp.route("/student/<int:student_id>/scores", methods=['GET', 'POST'])
@login_required
def manage_scores(student_id):
    """Trang nhập/sửa điểm cho một thiếu nhi."""
    student = Student.query.get_or_404(student_id)

    # Kiểm tra quyền: Admin hoặc Huynh trưởng của lớp, Guest không được sửa điểm
    if current_user.role == 'guest':
        flash('Bạn không có quyền sửa điểm thiếu nhi.', 'danger')
        abort(403)
    if not current_user.is_admin():
        if not student.class_id:
            flash('Thiếu nhi chưa được xếp lớp.', 'danger')
            abort(403)
        is_assigned = current_user.assigned_classes.filter_by(id=student.class_id).first()
        if not is_assigned:
            flash('Bạn không có quyền quản lý điểm của thiếu nhi này.', 'danger')
            abort(403)

    form = ScoreForm(obj=student)

    if form.validate_on_submit():
        student.diem_mieng = form.diem_mieng.data
        student.diem_giua_ki_1 = form.diem_giua_ki_1.data
        student.diem_cuoi_ki_1 = form.diem_cuoi_ki_1.data
        student.diem_giua_ki_2 = form.diem_giua_ki_2.data
        student.diem_cuoi_ki_2 = form.diem_cuoi_ki_2.data
        db.session.commit()
        flash('Điểm của thiếu nhi đã được cập nhật!', 'success')
        if current_user.is_admin():
            return redirect(url_for('student.list_students'))
        else:
            return redirect(url_for('main_routes.view_class', class_id=student.class_id))

    return render_template('student_scores.html', title=f'Quản lý điểm - {student.full_name}', form=form, student=student)

@student_bp.route("/students/export/<file_type>") # Route mới nhận file_type
@login_required
def export_students(file_type):
    """Xuất danh sách thiếu nhi ra file Excel (.xlsx) hoặc Word (.docx)."""

    # Lấy lại logic lọc/tìm kiếm từ URL parameters
    search_term = request.args.get('search_term', None)
    class_id_filter = request.args.get('class_filter', None)

    query = Student.query.join(Class).order_by(Class.name, Student.full_name)

    if search_term:
        search_pattern = f"%{search_term}%"
        query = query.filter(or_(
            Student.full_name.ilike(search_pattern),
            Student.ten_thanh.ilike(search_pattern)
        ))

    if class_id_filter:
        try:
            query = query.filter(Student.class_id == int(class_id_filter))
        except (ValueError, TypeError): # Bắt cả TypeError nếu class_id_filter là None
            pass 

    students = query.all()

    if not students:
        flash('Không có thiếu nhi nào để xuất file.', 'info')
        return redirect(url_for('student.list_students', search_term=search_term, class_filter=class_id_filter))

    # --- Tạo tên file ---
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    base_filename = f"danh_sach_sinh_vien_{timestamp}"

    # --- Xử lý EXCEL (.xlsx) ---
    if file_type == 'xlsx':
        # Chuyển đổi danh sách student objects thành list of dictionaries
        student_data = []
        for student in students:
            student_data.append({
                'Tên Thánh': student.ten_thanh or '',
                'Họ và tên': student.full_name,
                'Ngày sinh': student.date_of_birth.strftime('%d-%m-%Y'),
                'Giới tính': student.gender,
                'Lớp': student.class_assigned.name if student.class_assigned else '',
                'Họ tên Bố': student.ho_ten_bo or '',
                'Họ tên Mẹ': student.ho_ten_me or '',
                'SĐT Phụ huynh': student.sdt_phu_huynh or '',
                'Điểm miệng': student.diem_mieng or 0,
                'Điểm giữa kì 1': student.diem_giua_ki_1 or 0,
                'Điểm cuối kì 1': student.diem_cuoi_ki_1 or 0,
                'Điểm giữa kì 2': student.diem_giua_ki_2 or 0,
                'Điểm cuối kì 2': student.diem_cuoi_ki_2 or 0,
                'Điểm tổng': student.diem_tong
            })

        # Tạo DataFrame từ list of dictionaries
        df = pd.DataFrame(student_data)

        # Tạo file Excel trong bộ nhớ (BytesIO)
        output = BytesIO()
        writer = pd.ExcelWriter(output, engine='openpyxl')
        df.to_excel(writer, index=False, sheet_name='DanhSachSinhVien')
        # writer.save() # Không cần save với pandas mới
        writer.close() # Thay bằng close()
        output.seek(0) # Đưa con trỏ về đầu stream

        filename = f"{base_filename}.xlsx"
        mimetype = 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'

    # --- Xử lý WORD (.docx) ---
    elif file_type == 'docx':
        document = Document()
        document.add_heading('Danh sách thiếu nhi', 0)

        # Thêm bảng
        table = document.add_table(rows=1, cols=14) # 1 hàng tiêu đề, 14 cột
        table.style = 'Table Grid' # Kiểu bảng có kẻ ô
        table.autofit = False # Tắt tự động vừa cột

        # Set độ rộng cột (tùy chỉnh theo ý bạn)
        widths = (Inches(0.8), Inches(2.0), Inches(1.0), Inches(0.8), Inches(1.0), Inches(1.5), Inches(1.5), Inches(1.0), Inches(0.8), Inches(0.8), Inches(0.8), Inches(0.8), Inches(0.8), Inches(0.8))
        hdr_cells = table.rows[0].cells
        header_texts = ['Tên Thánh', 'Họ và tên', 'Ngày sinh', 'Giới tính', 'Lớp', 'Họ tên Bố', 'Họ tên Mẹ', 'SĐT Phụ huynh', 'Điểm miệng', 'Giữa kì 1', 'Cuối kì 1', 'Giữa kì 2', 'Cuối kì 2', 'Điểm tổng']
        for i, text in enumerate(header_texts):
             hdr_cells[i].text = text
             hdr_cells[i].width = widths[i]

        # Thêm dữ liệu thiếu nhi
        for student in students:
            row_cells = table.add_row().cells
            row_cells[0].text = student.ten_thanh or ''
            row_cells[1].text = student.full_name
            row_cells[2].text = student.date_of_birth.strftime('%d-%m-%Y')
            row_cells[3].text = student.gender
            row_cells[4].text = student.class_assigned.name if student.class_assigned else ''
            row_cells[5].text = student.ho_ten_bo or ''
            row_cells[6].text = student.ho_ten_me or ''
            row_cells[7].text = student.sdt_phu_huynh or ''
            row_cells[8].text = str(student.diem_mieng or 0)
            row_cells[9].text = str(student.diem_giua_ki_1 or 0)
            row_cells[10].text = str(student.diem_cuoi_ki_1 or 0)
            row_cells[11].text = str(student.diem_giua_ki_2 or 0)
            row_cells[12].text = str(student.diem_cuoi_ki_2 or 0)
            row_cells[13].text = str(student.diem_tong)
            # Căn chỉnh độ rộng cho các ô dữ liệu (tùy chọn)
            for i, width in enumerate(widths):
                row_cells[i].width = width

        # Tạo file Word trong bộ nhớ (BytesIO)
        output = BytesIO()
        document.save(output)
        output.seek(0)

        filename = f"{base_filename}.docx"
        mimetype = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'

    # --- Định dạng không hợp lệ ---
    else:
        flash('Định dạng file không được hỗ trợ.', 'danger')
        return redirect(url_for('student.list_students', search_term=search_term, class_filter=class_id_filter))

    # --- Tạo Response để tải file ---
    response = make_response(output.getvalue())
    response.headers["Content-Disposition"] = f"attachment; filename=\"{filename}\"" # Thêm dấu ngoặc kép cho tên file
    response.headers["Content-type"] = mimetype
    return response

# lấy dữ liệu từ file
UPLOAD_FOLDER = 'uploads'
ALLOWED_EXTENSIONS = {'xlsx', 'csv'}


def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


@student_bp.route('/students/import', methods=['GET', 'POST'])
@login_required
def import_students():
    # Chỉ admin và leader mới được import từ file
    if current_user.role == 'guest':
        flash('Bạn không có quyền nhập dữ liệu từ file.', 'danger')
        return redirect(url_for('student.list_students'))
    if request.method == 'POST':

        # === Nút XEM TRƯỚC ===
        if 'preview' in request.form:
            file = request.files.get('file')
            if not file or file.filename == '':
                flash('⚠️ Vui lòng chọn file!', 'warning')
                return redirect(request.url)

            if allowed_file(file.filename):
                filename = secure_filename(file.filename)
                os.makedirs(UPLOAD_FOLDER, exist_ok=True)
                filepath = os.path.join(UPLOAD_FOLDER, filename)
                file.save(filepath)

                try:
                    if filename.endswith('.xlsx'):
                        df = pd.read_excel(filepath)
                    else:
                        df = pd.read_csv(filepath)

                    # Lưu đường dẫn file vào session để dùng khi confirm
                    session['import_file'] = filepath

                    # Xem trước 5 dòng đầu tiên
                    preview_data = df.head(5).to_html(classes="table table-bordered table-striped table-sm", index=False)
                    return render_template('import_students.html', preview=preview_data)

                except Exception as e:
                    flash(f'❌ Lỗi khi đọc file: {e}', 'danger')
                    return redirect(request.url)
            else:
                flash('❌ Chỉ hỗ trợ file .xlsx hoặc .csv!', 'danger')
                return redirect(request.url)

        # === Nút XÁC NHẬN NHẬP ===
        elif 'confirm' in request.form:
            filepath = session.get('import_file')
            if not filepath or not os.path.exists(filepath):
                flash('❌ Không tìm thấy file để nhập. Vui lòng tải lại!', 'danger')
                return redirect(request.url)

            try:
                # Đọc lại file
                if filepath.endswith('.xlsx'):
                    df = pd.read_excel(filepath)
                else:
                    df = pd.read_csv(filepath)

                # Chuẩn hóa tên cột
                df.columns = df.columns.str.strip().str.lower()

                added_count = 0

                for _, row in df.iterrows():
                    ten_thanh = str(row.get("tên thánh", "")).strip()
                    full_name = str(row.get("họ và tên", "")).strip()
                    date_of_birth_str = str(row.get("ngày sinh", "")).strip()
                    gender = str(row.get("giới tính", "")).strip()
                    class_name = str(row.get("lớp", "")).strip()
                    ho_ten_bo = str(row.get("họ tên bố", "")).strip()
                    ho_ten_me = str(row.get("họ tên mẹ", "")).strip()
                    sdt_phu_huynh = str(row.get("sđt phụ huynh", "")).strip()

                    # Bỏ qua dòng thiếu dữ liệu chính
                    if not full_name:
                        continue

                    # Chuyển đổi ngày sinh
                    try:
                        date_of_birth = datetime.strptime(date_of_birth_str, "%d-%m-%Y").date()
                    except:
                        try:
                            date_of_birth = datetime.strptime(date_of_birth_str, "%d/%m/%Y").date()
                        except:
                            date_of_birth = None

                    # Tìm lớp trong DB
                    class_obj = Class.query.filter_by(name=class_name).first()

                    # Kiểm tra trùng tên (vì không có mã SV nữa)
                    if Student.query.filter_by(full_name=full_name, date_of_birth=date_of_birth).first():
                        continue

                    # Tạo thiếu nhi mới
                    new_student = Student(
                        ten_thanh=ten_thanh if ten_thanh else None,
                        full_name=full_name,
                        date_of_birth=date_of_birth,
                        gender=gender,
                        ho_ten_bo=ho_ten_bo if ho_ten_bo else None,
                        ho_ten_me=ho_ten_me if ho_ten_me else None,
                        sdt_phu_huynh=sdt_phu_huynh if sdt_phu_huynh else None,
                        class_id=class_obj.id if class_obj else None
                    )

                    db.session.add(new_student)
                    added_count += 1

                db.session.commit()
                flash(f"✅ Đã nhập thành công {added_count} thiếu nhi!", "success")

            except Exception as e:
                db.session.rollback()
                flash(f'❌ Lỗi khi xử lý file: {e}', 'danger')
                return redirect(request.url)

    # GET request hoặc render lại
    return render_template('import_students.html')
